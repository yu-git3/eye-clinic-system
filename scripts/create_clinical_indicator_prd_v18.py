from copy import deepcopy
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "02_需求文档_PRD" / "眼科专科系统_临床指标定义功能需求说明_PRD_V1.7.docx"
TARGET = ROOT / "02_需求文档_PRD" / "眼科专科系统_临床指标定义功能需求说明_PRD_V1.8.docx"


def replace_text(paragraph, old, new):
    if old not in paragraph.text:
        return
    text = paragraph.text.replace(old, new)
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


doc = Document(SOURCE)

# 文档标识与页眉页脚
doc.tables[0].rows[1].cells[1].text = "V1.8"
doc.tables[0].rows[1].cells[2].text = "2026-08-20"
for section in doc.sections:
    for container in (section.header, section.footer):
        for paragraph in container.paragraphs:
            replace_text(paragraph, "V1.7", "V1.8")
            replace_text(paragraph, "2026-08-18", "2026-08-20")
        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_text(paragraph, "V1.7", "V1.8")
                        replace_text(paragraph, "2026-08-18", "2026-08-20")

# 修订记录：保留V1.7，并增加V1.8说明
revision = doc.tables[1]
revision.rows[-1].cells[3].text = "历史版本"
new_row = deepcopy(revision.rows[-1]._tr)
revision._tbl.append(new_row)
for cell, value in zip(
    revision.rows[-1].cells,
    [
        "V1.8",
        "2026-08-20",
        "调整枚举结果属性互斥规则：0-正常与1-异常可共存；2-无与0-正常、1-异常均互斥；明确配置禁选、保存校验及多选录入自动取消规则。",
        "当前版本",
    ],
):
    cell.text = value

doc.save(TARGET)
print(TARGET)
