from copy import deepcopy
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "02_需求文档_PRD" / "眼科专科系统_临床指标定义功能需求说明_PRD_V1.5.docx"
TARGET = ROOT / "02_需求文档_PRD" / "眼科专科系统_临床指标定义功能需求说明_PRD_V1.7.docx"


def set_cell(cell, text):
    cell.text = text


def replace_paragraph(paragraph, old, new):
    if old not in paragraph.text:
        return False
    full = paragraph.text.replace(old, new)
    if paragraph.runs:
        paragraph.runs[0].text = full
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(full)
    return True


def append_copied_row(table, values):
    row = deepcopy(table.rows[-1]._tr)
    table._tbl.append(row)
    cells = table.rows[-1].cells
    for cell, value in zip(cells, values):
        set_cell(cell, value)


def insert_copied_row_before(table, before_index, values):
    row = deepcopy(table.rows[before_index]._tr)
    table.rows[before_index]._tr.addprevious(row)
    cells = table.rows[before_index].cells
    for cell, value in zip(cells, values):
        set_cell(cell, value)


doc = Document(SOURCE)

for section in doc.sections:
    for container in (section.header, section.footer):
        for paragraph in container.paragraphs:
            replace_paragraph(paragraph, "V1.5", "V1.7")
            replace_paragraph(paragraph, "2026-08-10", "2026-08-18")
        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_paragraph(paragraph, "V1.5", "V1.7")
                        replace_paragraph(paragraph, "2026-08-10", "2026-08-18")

# 封面与修订记录
set_cell(doc.tables[0].rows[1].cells[1], "V1.7")
set_cell(doc.tables[0].rows[1].cells[2], "2026-08-18")
revision = doc.tables[1]
set_cell(revision.rows[-1].cells[3], "历史版本")
append_copied_row(revision, ["V1.6", "2026-08-11", "补充指标安全删除规则：仅未引用且无历史数据时允许删除，其余仅允许停用；增加二次确认、审计与服务端并发校验。", "历史版本"])
append_copied_row(revision, ["V1.7", "2026-08-18", "枚举值收敛为结果属性、是否默认、是否补充文本三项；补充文本固定200字符；指标编码支持英文句点分层。", "当前版本"])

# 基础字段和编辑规则
set_cell(doc.tables[6].rows[2].cells[3], "全局唯一；允许大写字母、数字、下划线和英文句点；句点用于业务层级，不可连续或位于首尾；创建后不可修改。")
set_cell(doc.tables[7].rows[6].cells[0], "枚举/多选枚举项、结果属性、默认项、补充文本标志、布尔显示名及外部映射")

# 枚举规则表：三个属性与原有字段同行，位于排序前
enum_table = doc.tables[10]
insert_copied_row_before(enum_table, 4, ["结果属性", "0-正常、1-异常、2-无；默认2-无。同时存在0和1时，录入组件生成正常/异常快捷录入。"])
insert_copied_row_before(enum_table, 5, ["是否默认", "是/否；默认否。同一指标最多一个已启用的默认项。"])
insert_copied_row_before(enum_table, 6, ["是否补充文本", "是/否；默认否。选中该枚举值后展示补充文本框，最大长度固定200字符。"])

for paragraph in doc.paragraphs:
    replace_paragraph(
        paragraph,
        "枚举值由本页面维护、业务表存储。枚举型只能选一项；多选枚举型可选零项或多项并按枚举编码集合存储，不拼接展示名称。",
        "枚举值由本页面维护、业务表存储。枚举型只能选一项；多选枚举型可选零项或多项并按枚举编码集合存储。每个枚举值同行维护结果属性、是否默认、是否补充文本；不维护互斥属性和补充文本长度，补充文本长度统一为200字符。",
    )
    replace_paragraph(
        paragraph,
        "已被模板引用的指标不可物理删除。",
        "仅未被模板引用且无历史临床数据的指标允许删除；其余情况仅允许停用。删除前需二次确认，服务端重新校验引用和历史数据并记录审计信息。",
    )

all_paragraphs = list(doc.paragraphs)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            all_paragraphs.extend(cell.paragraphs)

for paragraph in all_paragraphs:
    replace_paragraph(
        paragraph,
        "每个已选眼别必须维护一个外部映射字段；选择无眼别时只维护一个通用字段。",
        "每个已选眼别可选维护一个外部映射字段；选择无眼别时可选维护一个通用字段。全部留空时允许保存，未配置的结果不自动转换写入。",
    )
    replace_paragraph(
        paragraph,
        "保存时每个选中眼别均须完成来源所需关联或映射。",
        "保存时，护士采集须为每个选中眼别完成护理体征关联；医技检查的外部字段为选填。",
    )
    replace_paragraph(
        paragraph,
        "仅未被模板引用且无历史临床数据的指标允许删除；其余情况仅允许停用。删除前需二次确认，",
        "仅未被模板引用且无历史临床数据的指标允许删除；其余情况不允许删除，只允许停用。删除后不可恢复，删除前需二次确认，",
    )

# 校验、数据与验收
set_cell(doc.tables[13].rows[2].cells[2], "编码必须全局唯一；仅允许大写字母、数字、下划线和英文句点；句点不可连续或位于首尾。")
set_cell(doc.tables[13].rows[8].cells[2], "校验枚举编码、名称、启用项、结果属性0/1/2、唯一默认项及补充文本标志。")
set_cell(doc.tables[10].rows[3].cells[1], "选填；医技接口转换使用；同一指标内已填写的编码不可重复。")
set_cell(doc.tables[11].rows[2].cells[1], "医技来源时选填；若维护则真、假映射需同时填写且不得重复。")
set_cell(doc.tables[13].rows[7].cells[0], "护理关联")
set_cell(doc.tables[13].rows[7].cells[1], "护士采集已选眼别缺少护理体征")
set_cell(doc.tables[13].rows[7].cells[2], "请按眼别关联护理体征。医技外部字段选填，已填写时校验格式和重复。")
set_cell(doc.tables[13].rows[9].cells[1], "显示名为空，或真/假外部映射只填一项/重复")
set_cell(doc.tables[13].rows[9].cells[2], "显示名必填；医技映射选填，填写时须成对且不得重复。")
set_cell(doc.tables[14].rows[6].cells[1], "指标ID、枚举编码、名称、外部映射、结果属性、默认标志、补充文本标志、排序、状态")
set_cell(doc.tables[18].rows[7].cells[1], "医技检查为每个选中眼别提供可选的外部映射字段；未配置时允许保存，但不自动转换写入。")
set_cell(doc.tables[18].rows[9].cells[1], "枚举由页面维护并存业务表；同行维护结果属性0/1/2、是否默认、是否补充文本；补充文本固定200字符。")
set_cell(doc.tables[18].rows[10].cells[1], "指标编码支持英文句点分层；编码、范围、长度、眼别、护理关联、医技映射及枚举校验有效。")
set_cell(doc.tables[18].rows[12].cells[1], "仅未引用且无历史数据的指标允许删除；其余仅允许停用；删除需二次确认、服务端并发校验和审计记录。")

doc.save(TARGET)
print(TARGET)
